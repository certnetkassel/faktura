import base64
import io
import json
import os
import shutil
import smtplib
import sqlite3
import subprocess
import tempfile
import time
from datetime import datetime, date, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from functools import wraps

import requests
from flask import (Flask, render_template, request, redirect, url_for,
                   flash, session, send_file, jsonify)
from werkzeug.security import generate_password_hash, check_password_hash

from werkzeug.utils import secure_filename

from config import DATABASE, SECRET_KEY, UPLOAD_FOLDER, OUTPUT_FOLDER
from database import get_db, init_db, migrate_db, EMAIL_TEMPLATE_DEFAULTS

LOGO_FOLDER = os.path.join(os.path.dirname(__file__), 'static', 'logos')

app = Flask(__name__)
app.secret_key = SECRET_KEY

# Unter Gunicorn wird der __main__-Block nicht ausgeführt, deshalb hier:
# fehlende Spalten in bestehenden Datenbanken nachziehen.
try:
    migrate_db()
except Exception as e:  # z.B. beim allerersten Start ohne Datenbank
    print(f"Migration übersprungen: {e}")

# ── Helpers ──

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('user_id'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    """Nur für Administratoren. Nicht-Admins landen auf dem Dashboard."""
    @wraps(f)
    def decorated(*args, **kwargs):
        u = current_user()
        if not u:
            return redirect(url_for('login'))
        if not u['is_admin']:
            flash('Dafür fehlen dir die Rechte (nur Administratoren).', 'error')
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated


def current_user():
    """Der aktuell angemeldete Benutzer als Row oder None."""
    uid = session.get('user_id')
    if not uid:
        return None
    db = get_db()
    u = db.execute("SELECT * FROM users WHERE id=?", [uid]).fetchone()
    db.close()
    return u


def get_settings():
    db = get_db()
    s = db.execute("SELECT * FROM settings WHERE id=1").fetchone()
    db.close()
    return s


def generate_doc_nr(prefix, seq_nr):
    """Generate document number in format PREFIX-YYMM### (e.g. RE-2603001)."""
    now = date.today()
    yymm = now.strftime('%y%m')
    return f"{prefix}-{yymm}{seq_nr:03d}"


def fmt_date(d):
    if not d:
        return ''
    try:
        return datetime.strptime(d, '%Y-%m-%d').strftime('%d.%m.%Y')
    except (ValueError, TypeError):
        return d


app.jinja_env.filters['fmt_date'] = fmt_date


# ── Word-Platzhalter-Ersetzung (robust) ──

def _replace_in_paragraph(paragraph, replacements):
    """Ersetzt Platzhalter in einem Absatz – auch wenn Word den Platzhalter
    über mehrere Runs zerrissen hat (z.B. '{{rechnung' + '_nr}}').

    Word teilt Platzhalter beim Bearbeiten oft auf mehrere Runs auf. Eine
    Ersetzung pro Run (key in run.text) scheitert dann. Daher wird der
    komplette Absatztext zusammengesetzt, ersetzt und – nur wenn sich etwas
    geändert hat – in den ersten Run geschrieben; die übrigen Runs werden
    geleert. Die Formatierung des ersten Runs bleibt erhalten.
    """
    if '{{' not in paragraph.text:
        return
    original = paragraph.text
    new_text = original
    for key, value in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, value if value is not None else '')
    if new_text == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.text = new_text


def _replace_in_tables(tables, replacements):
    """Ersetzt Platzhalter in allen Tabellenzellen, rekursiv (verschachtelte Tabellen)."""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
                if cell.tables:
                    _replace_in_tables(cell.tables, replacements)


def replace_placeholders(doc, replacements):
    """Ersetzt alle Platzhalter im gesamten Dokument: Body, Tabellen sowie
    Kopf- und Fußzeilen aller Abschnitte (inkl. First-Page/Even-Page)."""
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    _replace_in_tables(doc.tables, replacements)
    for section in doc.sections:
        for hf in (section.header, section.first_page_header, section.even_page_header,
                   section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in hf.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            _replace_in_tables(hf.tables, replacements)


LOGO_PLACEHOLDER = '{{logo}}'
LOGO_WIDTH_CM = 4.0  # Standardbreite des Logos im Dokument (Fallback)
LOGO_WIDTH_MIN_CM = 1.0
LOGO_WIDTH_MAX_CM = 10.0
SIDEBAR_LOGO_PX = 200  # Standardbreite des Logos in der Seitenleiste (Fallback)
SIDEBAR_LOGO_MIN_PX = 60
SIDEBAR_LOGO_MAX_PX = 240  # mehr als die Seitenleiste breit ist wäre sinnlos


def clamp(value, minimum, maximum, default):
    """Wandelt value in eine Zahl und begrenzt sie auf [minimum, maximum]."""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return default
    return max(minimum, min(maximum, number))


def get_logo_width_cm(s):
    """Logo-Breite für Dokumente aus den Einstellungen (in cm)."""
    try:
        value = s['logo_width_cm']
    except (KeyError, IndexError, TypeError):
        value = None
    return clamp(value, LOGO_WIDTH_MIN_CM, LOGO_WIDTH_MAX_CM, LOGO_WIDTH_CM)


def get_logo_sidebar_px(s):
    """Logo-Breite in der Seitenleiste aus den Einstellungen (in Pixeln)."""
    try:
        value = s['logo_sidebar_px']
    except (KeyError, IndexError, TypeError):
        value = None
    return int(clamp(value, SIDEBAR_LOGO_MIN_PX, SIDEBAR_LOGO_MAX_PX, SIDEBAR_LOGO_PX))


def _insert_logo_in_paragraph(paragraph, image_path, width):
    """Ersetzt den Marker {{logo}} in einem Absatz durch ein Inline-Bild.
    Funktioniert auch run-übergreifend (Word zerreißt Marker beim Bearbeiten).
    Reihenfolge Text-vor-Marker / Bild / Text-nach-Marker bleibt erhalten."""
    if LOGO_PLACEHOLDER not in paragraph.text:
        return
    if not paragraph.runs:
        return
    before, _, after = paragraph.text.partition(LOGO_PLACEHOLDER)

    # Häufiges Vorlagen-Muster: der Marker wurde mit einem Tab nach rechts
    # geschoben (z.B. "\t{{logo}}"). Ein Inline-Bild richtet sich aber NICHT am
    # Tabstopp aus – es startet an der Tab-Position und läuft über den rechten
    # Seitenrand hinaus. Steht vor dem Marker nur Tab/Leerraum, entfernen wir den
    # Tab und setzen den Absatz rechtsbündig: dann sitzt die rechte Bildkante
    # sauber am rechten Rand.
    right_align = before.strip() == '' and '\t' in before
    if right_align:
        before = before.replace('\t', '')

    run = paragraph.runs[0]
    run.text = before
    for r in paragraph.runs[1:]:
        r.text = ''
    run.add_picture(image_path, width=width)
    if after:
        run.add_text(after)

    if right_align:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _insert_logo_in_tables(tables, image_path, width):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _insert_logo_in_paragraph(paragraph, image_path, width)
                if cell.tables:
                    _insert_logo_in_tables(cell.tables, image_path, width)


def insert_logo(doc, image_path, width_cm=LOGO_WIDTH_CM):
    """Fügt das Logo überall dort ein, wo der Marker {{logo}} steht –
    Body, Tabellen sowie Kopf-/Fußzeilen aller Abschnitte."""
    from docx.shared import Cm
    width = Cm(width_cm)
    for paragraph in doc.paragraphs:
        _insert_logo_in_paragraph(paragraph, image_path, width)
    _insert_logo_in_tables(doc.tables, image_path, width)
    for section in doc.sections:
        for hf in (section.header, section.first_page_header, section.even_page_header,
                   section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in hf.paragraphs:
                _insert_logo_in_paragraph(paragraph, image_path, width)
            _insert_logo_in_tables(hf.tables, image_path, width)


def get_logo_path(s, prefer='light'):
    """Liefert den Dateipfad des hinterlegten Logos (Dokumente: helle Variante
    für weißen Hintergrund), oder None wenn keins existiert."""
    order = ['logo_light', 'logo_dark'] if prefer == 'light' else ['logo_dark', 'logo_light']
    for field in order:
        name = s[field] if s[field] else ''
        if name:
            path = os.path.join(LOGO_FOLDER, name)
            if os.path.exists(path):
                return path
    return None


def set_cell_text(cell, text):
    """Schreibt Text in eine Zelle, ohne ihre Formatierung zu verlieren.
    `cell.text = ...` würde Absatzausrichtung und Schriftformat verwerfen –
    deshalb den ersten Run weiterverwenden und den Rest entfernen."""
    paragraph = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def freeze_table_layout(table):
    """Tabellenbreiten festnageln. Ohne das rechnet Word/LibreOffice die Spalten
    nach Inhalt neu (Autofit): die schmale Pos.-Spalte wird breit, Menge,
    Einheit, Einzelpreis und Gesamt werden gequetscht. Das tblGrid der Vorlage
    weicht dabei von den tatsächlichen Zellbreiten ab, deshalb wird es aus der
    Kopfzeile neu aufgebaut."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    if not table.rows:
        return
    widths = [cell.width.twips if cell.width else None for cell in table.rows[0].cells]
    grid = table._tbl.find(qn('w:tblGrid'))
    if grid is not None and all(w for w in widths):
        for col, width in zip(grid.findall(qn('w:gridCol')), widths):
            col.set(qn('w:w'), str(width))


def fill_positions(doc, items):
    """Ersetzt den {{positionen}}-Marker durch je eine Tabellenzeile pro Position.
    Jede Zeile ist eine Kopie der Markerzeile – so bleiben Spaltenbreiten,
    Ausrichtung (z.B. rechtsbündige Beträge), Schrift und Rahmen der Vorlage
    erhalten. Die Markerzeile selbst wird anschließend entfernt (sonst bliebe an
    ihrer Stelle eine leere Zeile stehen)."""
    import copy

    from docx.table import _Row

    for table in doc.tables:
        marker_row = None
        for row in table.rows:
            if any('{{positionen}}' in cell.text for cell in row.cells):
                marker_row = row
                break
        if marker_row is None:
            continue
        for item in items:
            new_tr = copy.deepcopy(marker_row._tr)
            marker_row._tr.addprevious(new_tr)
            cells = _Row(new_tr, table).cells
            if len(cells) >= 5:
                set_cell_text(cells[0], str(item['position']))
                set_cell_text(cells[1], item['description'])
                set_cell_text(cells[2], str(item['quantity']))
                set_cell_text(cells[3], item['unit'])
                set_cell_text(cells[4], f"{item['price']:.2f} €")
                if len(cells) >= 6:
                    set_cell_text(cells[5], f"{item['total']:.2f} €")
        # Markerzeile entfernen
        marker_row._tr.getparent().remove(marker_row._tr)
        freeze_table_layout(table)
        break  # nur eine Positionstabelle


@app.context_processor
def inject_settings():
    """Make settings available in all templates (for logo display)."""
    try:
        s = get_settings()
        return {'settings': s, 'logo_sidebar_px': get_logo_sidebar_px(s),
                'current_user': current_user()}
    except Exception:
        return {'settings': {}, 'logo_sidebar_px': SIDEBAR_LOGO_PX, 'current_user': None}


# ── Auth ──

@app.route('/login', methods=['GET', 'POST'])
def login():
    db = get_db()
    user_count = db.execute("SELECT COUNT(*) FROM users").fetchone()[0]

    # Erststart (frische Datenbank ohne Benutzer): ersten Admin anlegen.
    if user_count == 0:
        if request.method == 'POST':
            pw = request.form.get('password', '')
            if len(pw) < 4:
                flash('Passwort muss mindestens 4 Zeichen haben.', 'error')
                db.close()
                return render_template('login.html', first_run=True)
            db.execute(
                "INSERT INTO users (email, password_hash, first_name, last_name, is_admin) "
                "VALUES (?, ?, ?, ?, 1)",
                ['dirk@dirkhildebrand.de', generate_password_hash(pw), 'Dirk', 'Hildebrand'])
            db.commit()
            new_id = db.execute("SELECT id FROM users WHERE email=?",
                                ['dirk@dirkhildebrand.de']).fetchone()['id']
            db.close()
            session['user_id'] = new_id
            flash('Startbenutzer dirk@dirkhildebrand.de angelegt. Willkommen!', 'success')
            return redirect(url_for('dashboard'))
        db.close()
        return render_template('login.html', first_run=True)

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        pw = request.form.get('password', '')
        user = db.execute("SELECT * FROM users WHERE email=?", [email]).fetchone()
        db.close()
        if user and check_password_hash(user['password_hash'], pw):
            session['user_id'] = user['id']
            return redirect(url_for('dashboard'))
        flash('E-Mail oder Passwort falsch.', 'error')
        return render_template('login.html')

    db.close()
    return render_template('login.html')


@app.route('/logout')
def logout():
    session.pop('user_id', None)
    session.pop('logged_in', None)  # Altlast aus dem Einzel-Passwort-Login
    return redirect(url_for('login'))


# ── Dashboard ──

@app.route('/')
@login_required
def dashboard():
    db = get_db()
    today = date.today().isoformat()
    month_start = date.today().replace(day=1).isoformat()

    stats = {
        'customers': db.execute("SELECT COUNT(*) FROM customers").fetchone()[0],
        'open_invoices': db.execute(
            "SELECT COUNT(*) FROM invoices WHERE status IN ('Gesendet','Überfällig')").fetchone()[0],
        'overdue_invoices': db.execute(
            "SELECT COUNT(*) FROM invoices WHERE status='Überfällig' OR (status='Gesendet' AND due_date < ?)",
            [today]).fetchone()[0],
        # Offen im Monat: gestellt (Gesendet/Überfällig), aber noch nicht bezahlt
        'month_open': db.execute(
            "SELECT COALESCE(SUM(total),0) FROM invoices "
            "WHERE status IN ('Gesendet','Überfällig') AND date >= ?",
            [month_start]).fetchone()[0],
        # Bezahlt im Monat: tatsächlich eingegangenes Geld
        'month_paid': db.execute(
            "SELECT COALESCE(SUM(total),0) FROM invoices WHERE status='Bezahlt' AND date >= ?",
            [month_start]).fetchone()[0],
    }

    # Update overdue status
    db.execute(
        "UPDATE invoices SET status='Überfällig' WHERE status='Gesendet' AND due_date < ?",
        [today])
    db.commit()

    open_invoices = db.execute('''
        SELECT i.*, c.last_name, c.company FROM invoices i
        JOIN customers c ON i.customer_id=c.id
        WHERE i.status IN ('Gesendet','Überfällig')
        ORDER BY i.due_date ASC LIMIT 10
    ''').fetchall()

    # Recent items
    recent_items = []
    for inv in db.execute(
            "SELECT i.*, c.last_name FROM invoices i JOIN customers c ON i.customer_id=c.id ORDER BY i.created_at DESC LIMIT 5").fetchall():
        recent_items.append({'type': 'Rechnung', 'badge': 'blue', 'nr': inv['invoice_nr'],
                             'customer': inv['last_name'], 'date': fmt_date(inv['date'])})
    for off in db.execute(
            "SELECT o.*, c.last_name FROM offers o JOIN customers c ON o.customer_id=c.id ORDER BY o.created_at DESC LIMIT 5").fetchall():
        recent_items.append({'type': 'Angebot', 'badge': 'green', 'nr': off['offer_nr'],
                             'customer': off['last_name'], 'date': fmt_date(off['date'])})
    recent_items.sort(key=lambda x: x['date'], reverse=True)

    db.close()
    return render_template('dashboard.html', stats=stats, open_invoices=open_invoices,
                           recent_items=recent_items[:10])


# ── Settings ──

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if request.method == 'POST':
        db = get_db()
        fields = ['company_name', 'owner_name', 'street', 'zip', 'city', 'phone', 'email',
                  'tax_number', 'bank_name', 'iban', 'bic', 'smtp_host', 'smtp_port',
                  'smtp_user', 'smtp_pass', 'smtp_from', 'kleinunternehmer_text',
                  'invoice_prefix', 'offer_prefix', 'credit_prefix',
                  'next_invoice_nr', 'next_offer_nr', 'next_credit_nr',
                  'graph_tenant_id', 'graph_client_id', 'graph_sender']
        vals = {}
        for f in fields:
            vals[f] = request.form.get(f, '')
        sets = ', '.join(f"{f}=?" for f in fields)
        db.execute(f"UPDATE settings SET {sets} WHERE id=1", list(vals.values()))

        # Versandverfahren und Graph-Optionen
        method = 'graph' if request.form.get('mail_method') == 'graph' else 'smtp'
        db.execute("UPDATE settings SET mail_method=?, graph_save_sent=? WHERE id=1",
                   [method, 1 if request.form.get('graph_save_sent') else 0])

        # Client-Secret nur überschreiben, wenn ein neues eingegeben wurde
        new_secret = request.form.get('graph_client_secret', '')
        if new_secret:
            db.execute("UPDATE settings SET graph_client_secret=? WHERE id=1", [new_secret])

        # Logo-Größen (Schieberegler) – auf sinnvolle Bereiche begrenzen
        width_cm = clamp(request.form.get('logo_width_cm'),
                         LOGO_WIDTH_MIN_CM, LOGO_WIDTH_MAX_CM, LOGO_WIDTH_CM)
        sidebar_px = clamp(request.form.get('logo_sidebar_px'),
                           SIDEBAR_LOGO_MIN_PX, SIDEBAR_LOGO_MAX_PX, SIDEBAR_LOGO_PX)
        db.execute("UPDATE settings SET logo_width_cm=?, logo_sidebar_px=? WHERE id=1",
                   [round(width_cm, 1), int(sidebar_px)])

        # Handle logo uploads
        os.makedirs(LOGO_FOLDER, exist_ok=True)
        for logo_field in ['logo_dark', 'logo_light']:
            file = request.files.get(logo_field)
            if file and file.filename:
                filename = secure_filename(f"{logo_field}_{file.filename}")
                try:
                    file.save(os.path.join(LOGO_FOLDER, filename))
                except OSError as e:
                    flash(f'Logo konnte nicht gespeichert werden: {e}', 'error')
                    continue
                db.execute(f"UPDATE settings SET {logo_field}=? WHERE id=1", [filename])

        new_pw = request.form.get('new_password', '')
        if new_pw:
            if len(new_pw) < 4:
                flash('Passwort muss mindestens 4 Zeichen haben.', 'error')
            else:
                db.execute("UPDATE users SET password_hash=? WHERE id=?",
                           [generate_password_hash(new_pw), session['user_id']])
        db.commit()
        db.close()
        flash('Einstellungen gespeichert.', 'success')
        return redirect(url_for('settings'))

    s = get_settings()
    return render_template('settings.html', s=s,
                           mail_method=get_mail_method(s),
                           graph_secret_set=bool(setting(s, 'graph_client_secret')))


# ── Benutzerverwaltung (nur Admins) ──

@app.route('/users')
@admin_required
def users():
    db = get_db()
    all_users = db.execute(
        "SELECT * FROM users ORDER BY is_admin DESC, last_name, first_name").fetchall()
    db.close()
    return render_template('users.html', users=all_users)


@app.route('/users/new', methods=['GET', 'POST'])
@admin_required
def user_new():
    if request.method == 'POST':
        error = _save_user(None)
        if error:
            flash(error, 'error')
            return render_template('user_form.html', user=None, form=request.form)
        flash('Benutzer angelegt.', 'success')
        return redirect(url_for('users'))
    return render_template('user_form.html', user=None, form={})


@app.route('/users/<int:user_id>/edit', methods=['GET', 'POST'])
@admin_required
def user_edit(user_id):
    db = get_db()
    user = db.execute("SELECT * FROM users WHERE id=?", [user_id]).fetchone()
    db.close()
    if not user:
        flash('Benutzer nicht gefunden.', 'error')
        return redirect(url_for('users'))
    if request.method == 'POST':
        error = _save_user(user_id)
        if error:
            flash(error, 'error')
            return render_template('user_form.html', user=user, form=request.form)
        flash('Benutzer gespeichert.', 'success')
        return redirect(url_for('users'))
    return render_template('user_form.html', user=user, form=user)


@app.route('/users/<int:user_id>/delete', methods=['POST'])
@admin_required
def user_delete(user_id):
    if user_id == session.get('user_id'):
        flash('Du kannst dich nicht selbst löschen.', 'error')
        return redirect(url_for('users'))
    db = get_db()
    target = db.execute("SELECT * FROM users WHERE id=?", [user_id]).fetchone()
    if not target:
        db.close()
        flash('Benutzer nicht gefunden.', 'error')
        return redirect(url_for('users'))
    # Letzten Administrator nicht löschen
    if target['is_admin']:
        admin_count = db.execute(
            "SELECT COUNT(*) FROM users WHERE is_admin=1").fetchone()[0]
        if admin_count <= 1:
            db.close()
            flash('Der letzte Administrator kann nicht gelöscht werden.', 'error')
            return redirect(url_for('users'))
    db.execute("DELETE FROM users WHERE id=?", [user_id])
    db.commit()
    db.close()
    flash('Benutzer gelöscht.', 'success')
    return redirect(url_for('users'))


def _save_user(user_id):
    """Legt einen Benutzer an oder aktualisiert ihn. Gibt eine Fehlermeldung
    (String) zurück oder None bei Erfolg. Erwartet ein offenes request.form."""
    email = request.form.get('email', '').strip().lower()
    first_name = request.form.get('first_name', '').strip()
    last_name = request.form.get('last_name', '').strip()
    pw = request.form.get('password', '')
    is_admin = 1 if request.form.get('is_admin') else 0

    if not email or '@' not in email:
        return 'Bitte eine gültige E-Mail-Adresse angeben.'
    if user_id is None and len(pw) < 4:
        return 'Passwort muss mindestens 4 Zeichen haben.'
    if pw and len(pw) < 4:
        return 'Passwort muss mindestens 4 Zeichen haben.'

    db = get_db()
    # E-Mail-Eindeutigkeit prüfen
    clash = db.execute("SELECT id FROM users WHERE email=? AND id IS NOT ?",
                       [email, user_id]).fetchone()
    if clash:
        db.close()
        return 'Diese E-Mail-Adresse ist bereits vergeben.'

    # Verhindern, dass sich der letzte Admin selbst die Admin-Rechte entzieht
    if user_id is not None and not is_admin:
        was_admin = db.execute("SELECT is_admin FROM users WHERE id=?",
                               [user_id]).fetchone()['is_admin']
        if was_admin:
            admin_count = db.execute(
                "SELECT COUNT(*) FROM users WHERE is_admin=1").fetchone()[0]
            if admin_count <= 1:
                db.close()
                return 'Der letzte Administrator muss Administrator bleiben.'

    if user_id is None:
        db.execute(
            "INSERT INTO users (email, password_hash, first_name, last_name, is_admin) "
            "VALUES (?, ?, ?, ?, ?)",
            [email, generate_password_hash(pw), first_name, last_name, is_admin])
    else:
        db.execute(
            "UPDATE users SET email=?, first_name=?, last_name=?, is_admin=? WHERE id=?",
            [email, first_name, last_name, is_admin, user_id])
        if pw:
            db.execute("UPDATE users SET password_hash=? WHERE id=?",
                       [generate_password_hash(pw), user_id])
    db.commit()
    db.close()
    return None


# ── Customers ──

@app.route('/customers')
@login_required
def customers():
    db = get_db()
    custs = db.execute("SELECT * FROM customers ORDER BY last_name").fetchall()
    db.close()
    return render_template('customers.html', customers=custs)


@app.route('/customers/new', methods=['GET', 'POST'])
@login_required
def customer_form():
    if request.method == 'POST':
        db = get_db()
        nr = request.form['customer_nr']
        db.execute('''INSERT INTO customers (customer_nr, company, salutation, first_name, last_name,
                      street, zip, city, email, phone, notes)
                      VALUES (?,?,?,?,?,?,?,?,?,?,?)''',
                   [nr, request.form.get('company',''), request.form.get('salutation',''),
                    request.form.get('first_name',''), request.form['last_name'],
                    request.form.get('street',''), request.form.get('zip',''),
                    request.form.get('city',''), request.form.get('email',''),
                    request.form.get('phone',''), request.form.get('notes','')])
        db.commit()
        db.close()
        flash('Kunde angelegt.', 'success')
        return redirect(url_for('customers'))
    # Generate next customer nr
    db = get_db()
    last = db.execute("SELECT customer_nr FROM customers ORDER BY id DESC LIMIT 1").fetchone()
    next_nr = 'K-0001'
    if last:
        try:
            num = int(last['customer_nr'].split('-')[1]) + 1
            next_nr = f"K-{num:04d}"
        except (ValueError, IndexError):
            pass
    db.close()
    return render_template('customer_form.html', c=None, next_nr=next_nr)


@app.route('/customers/<int:id>/edit', methods=['GET', 'POST'])
@login_required
def customer_edit(id):
    db = get_db()
    if request.method == 'POST':
        db.execute('''UPDATE customers SET customer_nr=?, company=?, salutation=?, first_name=?,
                      last_name=?, street=?, zip=?, city=?, email=?, phone=?, notes=? WHERE id=?''',
                   [request.form['customer_nr'], request.form.get('company',''),
                    request.form.get('salutation',''), request.form.get('first_name',''),
                    request.form['last_name'], request.form.get('street',''),
                    request.form.get('zip',''), request.form.get('city',''),
                    request.form.get('email',''), request.form.get('phone',''),
                    request.form.get('notes',''), id])
        db.commit()
        db.close()
        flash('Kunde aktualisiert.', 'success')
        return redirect(url_for('customers'))
    c = db.execute("SELECT * FROM customers WHERE id=?", [id]).fetchone()
    db.close()
    return render_template('customer_form.html', c=c, next_nr=None)


@app.route('/customers/<int:id>/delete', methods=['POST'])
@login_required
def customer_delete(id):
    db = get_db()
    db.execute("DELETE FROM customers WHERE id=?", [id])
    db.commit()
    db.close()
    flash('Kunde gelöscht.', 'success')
    return redirect(url_for('customers'))


# ── Articles ──

@app.route('/articles')
@login_required
def articles():
    db = get_db()
    arts = db.execute("SELECT * FROM articles ORDER BY name").fetchall()
    db.close()
    return render_template('articles.html', articles=arts)


@app.route('/articles/new', methods=['GET', 'POST'])
@login_required
def article_form():
    if request.method == 'POST':
        db = get_db()
        db.execute('''INSERT INTO articles (article_nr, name, description, unit, price)
                      VALUES (?,?,?,?,?)''',
                   [request.form['article_nr'], request.form['name'],
                    request.form.get('description',''), request.form.get('unit','Stunde'),
                    float(request.form.get('price', 0))])
        db.commit()
        db.close()
        flash('Artikel angelegt.', 'success')
        return redirect(url_for('articles'))
    db = get_db()
    last = db.execute("SELECT article_nr FROM articles ORDER BY id DESC LIMIT 1").fetchone()
    next_nr = 'A-0001'
    if last:
        try:
            num = int(last['article_nr'].split('-')[1]) + 1
            next_nr = f"A-{num:04d}"
        except (ValueError, IndexError):
            pass
    db.close()
    return render_template('article_form.html', a=None, next_nr=next_nr)


@app.route('/articles/<int:id>/edit', methods=['GET', 'POST'])
@login_required
def article_edit(id):
    db = get_db()
    if request.method == 'POST':
        db.execute('''UPDATE articles SET article_nr=?, name=?, description=?, unit=?, price=?
                      WHERE id=?''',
                   [request.form['article_nr'], request.form['name'],
                    request.form.get('description',''), request.form.get('unit','Stunde'),
                    float(request.form.get('price', 0)), id])
        db.commit()
        db.close()
        flash('Artikel aktualisiert.', 'success')
        return redirect(url_for('articles'))
    a = db.execute("SELECT * FROM articles WHERE id=?", [id]).fetchone()
    db.close()
    return render_template('article_form.html', a=a, next_nr=None)


@app.route('/articles/<int:id>/delete', methods=['POST'])
@login_required
def article_delete(id):
    db = get_db()
    db.execute("DELETE FROM articles WHERE id=?", [id])
    db.commit()
    db.close()
    flash('Artikel gelöscht.', 'success')
    return redirect(url_for('articles'))


# ── Offers ──

@app.route('/offers')
@login_required
def offers():
    db = get_db()
    offs = db.execute('''SELECT o.*, c.last_name, c.company FROM offers o
                         JOIN customers c ON o.customer_id=c.id ORDER BY o.date DESC''').fetchall()
    db.close()
    return render_template('offers.html', offers=offs)


@app.route('/offers/new', methods=['GET', 'POST'])
@login_required
def offer_form():
    db = get_db()
    if request.method == 'POST':
        s = get_settings()
        offer_nr = generate_doc_nr(s['offer_prefix'], s['next_offer_nr'])
        customer_id = int(request.form['customer_id'])
        offer_date = request.form['date']
        valid_until = request.form.get('valid_until', '')
        notes = request.form.get('notes', '')

        db.execute("INSERT INTO offers (offer_nr,customer_id,date,valid_until,status,notes) VALUES(?,?,?,?,?,?)",
                   [offer_nr, customer_id, offer_date, valid_until, 'Entwurf', notes])
        offer_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]

        descriptions = request.form.getlist('item_desc')
        quantities = request.form.getlist('item_qty')
        units = request.form.getlist('item_unit')
        prices = request.form.getlist('item_price')

        total = 0
        for i, desc in enumerate(descriptions):
            if not desc.strip():
                continue
            qty = float(quantities[i] or 1)
            price = float(prices[i] or 0)
            item_total = qty * price
            total += item_total
            db.execute('''INSERT INTO offer_items (offer_id,position,description,quantity,unit,price,total)
                          VALUES(?,?,?,?,?,?,?)''',
                       [offer_id, i+1, desc, qty, units[i], price, item_total])

        db.execute("UPDATE offers SET total=? WHERE id=?", [total, offer_id])
        db.execute("UPDATE settings SET next_offer_nr=next_offer_nr+1 WHERE id=1")
        db.commit()
        db.close()
        flash('Angebot erstellt.', 'success')
        return redirect(url_for('offers'))

    customers = db.execute("SELECT * FROM customers ORDER BY last_name").fetchall()
    articles_list = db.execute("SELECT * FROM articles ORDER BY name").fetchall()
    db.close()
    return render_template('offer_form.html', o=None, customers=customers, articles=articles_list, items=[])


@app.route('/offers/<int:id>/edit', methods=['GET', 'POST'])
@login_required
def offer_edit(id):
    db = get_db()
    if request.method == 'POST':
        customer_id = int(request.form['customer_id'])
        db.execute("UPDATE offers SET customer_id=?,date=?,valid_until=?,notes=? WHERE id=?",
                   [customer_id, request.form['date'], request.form.get('valid_until',''),
                    request.form.get('notes',''), id])
        db.execute("DELETE FROM offer_items WHERE offer_id=?", [id])

        descriptions = request.form.getlist('item_desc')
        quantities = request.form.getlist('item_qty')
        units = request.form.getlist('item_unit')
        prices = request.form.getlist('item_price')

        total = 0
        for i, desc in enumerate(descriptions):
            if not desc.strip():
                continue
            qty = float(quantities[i] or 1)
            price = float(prices[i] or 0)
            item_total = qty * price
            total += item_total
            db.execute('''INSERT INTO offer_items (offer_id,position,description,quantity,unit,price,total)
                          VALUES(?,?,?,?,?,?,?)''',
                       [id, i+1, desc, qty, units[i], price, item_total])

        db.execute("UPDATE offers SET total=? WHERE id=?", [total, id])
        db.commit()
        db.close()
        flash('Angebot aktualisiert.', 'success')
        return redirect(url_for('offers'))

    o = db.execute("SELECT * FROM offers WHERE id=?", [id]).fetchone()
    items = db.execute("SELECT * FROM offer_items WHERE offer_id=? ORDER BY position", [id]).fetchall()
    customers = db.execute("SELECT * FROM customers ORDER BY last_name").fetchall()
    articles_list = db.execute("SELECT * FROM articles ORDER BY name").fetchall()
    db.close()
    return render_template('offer_form.html', o=o, customers=customers, articles=articles_list, items=items)


@app.route('/offers/<int:id>/delete', methods=['POST'])
@login_required
def offer_delete(id):
    db = get_db()
    db.execute("DELETE FROM offers WHERE id=?", [id])
    db.commit()
    db.close()
    flash('Angebot gelöscht.', 'success')
    return redirect(url_for('offers'))


@app.route('/offers/<int:id>/status/<status>', methods=['POST'])
@login_required
def offer_status(id, status):
    db = get_db()
    db.execute("UPDATE offers SET status=? WHERE id=?", [status, id])
    db.commit()
    db.close()
    flash(f'Status auf "{status}" gesetzt.', 'success')
    return redirect(url_for('offers'))


@app.route('/offers/<int:id>/to-invoice', methods=['POST'])
@login_required
def offer_to_invoice(id):
    db = get_db()
    s = get_settings()
    offer = db.execute("SELECT * FROM offers WHERE id=?", [id]).fetchone()
    items = db.execute("SELECT * FROM offer_items WHERE offer_id=? ORDER BY position", [id]).fetchall()

    invoice_nr = generate_doc_nr(s['invoice_prefix'], s['next_invoice_nr'])
    today = date.today().isoformat()
    due = (date.today() + timedelta(days=14)).isoformat()

    db.execute('''INSERT INTO invoices (invoice_nr,customer_id,offer_id,date,due_date,status,notes,total)
                  VALUES(?,?,?,?,?,?,?,?)''',
               [invoice_nr, offer['customer_id'], id, today, due, 'Entwurf',
                offer['notes'], offer['total']])
    inv_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]

    for item in items:
        db.execute('''INSERT INTO invoice_items (invoice_id,position,description,quantity,unit,price,total)
                      VALUES(?,?,?,?,?,?,?)''',
                   [inv_id, item['position'], item['description'], item['quantity'],
                    item['unit'], item['price'], item['total']])

    db.execute("UPDATE offers SET status='Angenommen' WHERE id=?", [id])
    db.execute("UPDATE settings SET next_invoice_nr=next_invoice_nr+1 WHERE id=1")
    db.commit()
    db.close()
    flash(f'Rechnung {invoice_nr} aus Angebot erstellt.', 'success')
    return redirect(url_for('invoices'))


# ── Invoices ──

# Erlaubte Rechnungs-Status (Reihenfolge = Workflow, auch für das Status-Dropdown)
INVOICE_STATUSES = ['Entwurf', 'Gesendet', 'Bezahlt', 'Überfällig', 'Storniert']


@app.route('/invoices')
@login_required
def invoices():
    db = get_db()
    today = date.today().isoformat()
    db.execute("UPDATE invoices SET status='Überfällig' WHERE status='Gesendet' AND due_date < ?", [today])
    db.commit()
    invs = db.execute('''SELECT i.*, c.last_name, c.company FROM invoices i
                         JOIN customers c ON i.customer_id=c.id ORDER BY i.date DESC''').fetchall()
    db.close()
    return render_template('invoices.html', invoices=invs, invoice_statuses=INVOICE_STATUSES)


@app.route('/invoices/new', methods=['GET', 'POST'])
@login_required
def invoice_form():
    db = get_db()
    if request.method == 'POST':
        s = get_settings()
        invoice_nr = generate_doc_nr(s['invoice_prefix'], s['next_invoice_nr'])
        customer_id = int(request.form['customer_id'])
        inv_date = request.form['date']
        due_date = request.form['due_date']
        notes = request.form.get('notes', '')

        db.execute('''INSERT INTO invoices (invoice_nr,customer_id,date,due_date,status,notes)
                      VALUES(?,?,?,?,?,?)''',
                   [invoice_nr, customer_id, inv_date, due_date, 'Entwurf', notes])
        inv_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]

        descriptions = request.form.getlist('item_desc')
        quantities = request.form.getlist('item_qty')
        units = request.form.getlist('item_unit')
        prices = request.form.getlist('item_price')

        total = 0
        for i, desc in enumerate(descriptions):
            if not desc.strip():
                continue
            qty = float(quantities[i] or 1)
            price = float(prices[i] or 0)
            item_total = qty * price
            total += item_total
            db.execute('''INSERT INTO invoice_items (invoice_id,position,description,quantity,unit,price,total)
                          VALUES(?,?,?,?,?,?,?)''',
                       [inv_id, i+1, desc, qty, units[i], price, item_total])

        db.execute("UPDATE invoices SET total=? WHERE id=?", [total, inv_id])
        db.execute("UPDATE settings SET next_invoice_nr=next_invoice_nr+1 WHERE id=1")
        db.commit()
        db.close()
        flash('Rechnung erstellt.', 'success')
        return redirect(url_for('invoices'))

    customers = db.execute("SELECT * FROM customers ORDER BY last_name").fetchall()
    articles_list = db.execute("SELECT * FROM articles ORDER BY name").fetchall()
    db.close()
    today = date.today().isoformat()
    due = (date.today() + timedelta(days=14)).isoformat()
    return render_template('invoice_form.html', inv=None, customers=customers,
                           articles=articles_list, items=[], today=today, due=due)


@app.route('/invoices/<int:id>/edit', methods=['GET', 'POST'])
@login_required
def invoice_edit(id):
    db = get_db()
    if request.method == 'POST':
        db.execute("UPDATE invoices SET customer_id=?,date=?,due_date=?,notes=? WHERE id=?",
                   [int(request.form['customer_id']), request.form['date'],
                    request.form['due_date'], request.form.get('notes',''), id])
        db.execute("DELETE FROM invoice_items WHERE invoice_id=?", [id])

        descriptions = request.form.getlist('item_desc')
        quantities = request.form.getlist('item_qty')
        units = request.form.getlist('item_unit')
        prices = request.form.getlist('item_price')

        total = 0
        for i, desc in enumerate(descriptions):
            if not desc.strip():
                continue
            qty = float(quantities[i] or 1)
            price = float(prices[i] or 0)
            item_total = qty * price
            total += item_total
            db.execute('''INSERT INTO invoice_items (invoice_id,position,description,quantity,unit,price,total)
                          VALUES(?,?,?,?,?,?,?)''',
                       [id, i+1, desc, qty, units[i], price, item_total])

        db.execute("UPDATE invoices SET total=? WHERE id=?", [total, id])
        db.commit()
        db.close()
        flash('Rechnung aktualisiert.', 'success')
        return redirect(url_for('invoices'))

    inv = db.execute("SELECT * FROM invoices WHERE id=?", [id]).fetchone()
    items = db.execute("SELECT * FROM invoice_items WHERE invoice_id=? ORDER BY position", [id]).fetchall()
    customers = db.execute("SELECT * FROM customers ORDER BY last_name").fetchall()
    articles_list = db.execute("SELECT * FROM articles ORDER BY name").fetchall()
    db.close()
    return render_template('invoice_form.html', inv=inv, customers=customers,
                           articles=articles_list, items=items, today=None, due=None)


@app.route('/invoices/<int:id>/delete', methods=['POST'])
@login_required
def invoice_delete(id):
    db = get_db()
    db.execute("DELETE FROM invoices WHERE id=?", [id])
    db.commit()
    db.close()
    flash('Rechnung gelöscht.', 'success')
    return redirect(url_for('invoices'))


@app.route('/invoices/<int:id>/status', methods=['POST'])
@login_required
def invoice_status(id):
    status = request.form.get('status', '')
    if status not in INVOICE_STATUSES:
        flash('Ungültiger Status.', 'error')
        return redirect(url_for('invoices'))
    db = get_db()
    db.execute("UPDATE invoices SET status=? WHERE id=?", [status, id])
    db.commit()
    db.close()
    flash(f'Status auf „{status}" gesetzt.', 'success')
    return redirect(url_for('invoices'))


# ── Credits ──

@app.route('/credits')
@login_required
def credits_list():
    db = get_db()
    creds = db.execute('''SELECT cr.*, c.last_name, c.company FROM credits cr
                          JOIN customers c ON cr.customer_id=c.id ORDER BY cr.date DESC''').fetchall()
    db.close()
    return render_template('credits.html', credits=creds)


@app.route('/credits/new', methods=['GET', 'POST'])
@login_required
def credit_form():
    db = get_db()
    if request.method == 'POST':
        s = get_settings()
        credit_nr = generate_doc_nr(s['credit_prefix'], s['next_credit_nr'])
        customer_id = int(request.form['customer_id'])
        invoice_id = request.form.get('invoice_id') or None
        cr_date = request.form['date']
        notes = request.form.get('notes', '')

        db.execute('''INSERT INTO credits (credit_nr,customer_id,invoice_id,date,notes)
                      VALUES(?,?,?,?,?)''',
                   [credit_nr, customer_id, invoice_id, cr_date, notes])
        cr_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]

        descriptions = request.form.getlist('item_desc')
        quantities = request.form.getlist('item_qty')
        units = request.form.getlist('item_unit')
        prices = request.form.getlist('item_price')

        total = 0
        for i, desc in enumerate(descriptions):
            if not desc.strip():
                continue
            qty = float(quantities[i] or 1)
            price = float(prices[i] or 0)
            item_total = qty * price
            total += item_total
            db.execute('''INSERT INTO credit_items (credit_id,position,description,quantity,unit,price,total)
                          VALUES(?,?,?,?,?,?,?)''',
                       [cr_id, i+1, desc, qty, units[i], price, item_total])

        db.execute("UPDATE credits SET total=? WHERE id=?", [total, cr_id])
        db.execute("UPDATE settings SET next_credit_nr=next_credit_nr+1 WHERE id=1")
        db.commit()
        db.close()
        flash('Gutschrift erstellt.', 'success')
        return redirect(url_for('credits_list'))

    customers = db.execute("SELECT * FROM customers ORDER BY last_name").fetchall()
    articles_list = db.execute("SELECT * FROM articles ORDER BY name").fetchall()
    invoices_list = db.execute("SELECT id, invoice_nr FROM invoices ORDER BY date DESC").fetchall()
    db.close()
    return render_template('credit_form.html', cr=None, customers=customers,
                           articles=articles_list, invoices=invoices_list, items=[])


@app.route('/credits/<int:id>/delete', methods=['POST'])
@login_required
def credit_delete(id):
    db = get_db()
    db.execute("DELETE FROM credits WHERE id=?", [id])
    db.commit()
    db.close()
    flash('Gutschrift gelöscht.', 'success')
    return redirect(url_for('credits_list'))


# ── Reminders ──

@app.route('/reminders')
@login_required
def reminders():
    db = get_db()
    rems = db.execute('''SELECT r.*, i.invoice_nr, i.total as inv_total, c.last_name, c.company
                         FROM reminders r
                         JOIN invoices i ON r.invoice_id=i.id
                         JOIN customers c ON i.customer_id=c.id
                         ORDER BY r.date DESC''').fetchall()
    overdue = db.execute('''SELECT i.*, c.last_name, c.company FROM invoices i
                            JOIN customers c ON i.customer_id=c.id
                            WHERE i.status='Überfällig' ORDER BY i.due_date''').fetchall()
    db.close()
    return render_template('reminders.html', reminders=rems, overdue=overdue)


@app.route('/reminders/new/<int:invoice_id>', methods=['GET', 'POST'])
@login_required
def reminder_form(invoice_id):
    db = get_db()
    if request.method == 'POST':
        level = int(request.form.get('level', 1))
        rem_date = request.form['date']
        due_date = request.form['due_date']
        fee = float(request.form.get('fee', 0))
        notes = request.form.get('notes', '')

        db.execute('''INSERT INTO reminders (invoice_id,level,date,due_date,fee,notes)
                      VALUES(?,?,?,?,?,?)''',
                   [invoice_id, level, rem_date, due_date, fee, notes])
        db.commit()
        db.close()
        flash(f'{level}. Mahnung erstellt.', 'success')
        return redirect(url_for('reminders'))

    invoice = db.execute('''SELECT i.*, c.last_name, c.company FROM invoices i
                            JOIN customers c ON i.customer_id=c.id WHERE i.id=?''',
                         [invoice_id]).fetchone()
    last_reminder = db.execute(
        "SELECT MAX(level) as max_level FROM reminders WHERE invoice_id=?",
        [invoice_id]).fetchone()
    next_level = (last_reminder['max_level'] or 0) + 1
    db.close()
    today = date.today().isoformat()
    due = (date.today() + timedelta(days=7)).isoformat()
    return render_template('reminder_form.html', invoice=invoice, level=next_level,
                           today=today, due=due)


@app.route('/reminders/<int:id>/delete', methods=['POST'])
@login_required
def reminder_delete(id):
    db = get_db()
    db.execute("DELETE FROM reminders WHERE id=?", [id])
    db.commit()
    db.close()
    flash('Mahnung gelöscht.', 'success')
    return redirect(url_for('reminders'))


# ── Document Generation (Word) ──

class DocGenError(Exception):
    """Fehler bei der Dokumenterzeugung (z.B. fehlende Vorlage). Die Nachricht
    ist für den Anwender bestimmt (wird als Flash angezeigt)."""


def company_placeholders(s):
    """Platzhalter aus den Einstellungen. Gemeinsame Quelle für Word-Vorlagen
    und E-Mail-Vorlagen, damit beide dieselben Namen und Werte kennen."""
    return {
        '{{firma}}': setting(s, 'company_name'),
        '{{inhaber}}': setting(s, 'owner_name'),
        '{{firma_strasse}}': setting(s, 'street'),
        '{{firma_plz}}': setting(s, 'zip'),
        '{{firma_stadt}}': setting(s, 'city'),
        '{{firma_telefon}}': setting(s, 'phone'),
        '{{firma_email}}': setting(s, 'email'),
        '{{steuernummer}}': setting(s, 'tax_number'),
        '{{bank}}': setting(s, 'bank_name'),
        '{{iban}}': setting(s, 'iban'),
        '{{bic}}': setting(s, 'bic'),
        '{{kleinunternehmer}}': setting(s, 'kleinunternehmer_text'),
    }


def customer_placeholders(cust):
    """Platzhalter aus den Kundendaten – ebenfalls für beide Vorlagenarten."""
    return {
        '{{kunde_firma}}': cust['company'],
        '{{kunde_anrede}}': cust['salutation'],
        '{{kunde_vorname}}': cust['first_name'],
        '{{kunde_nachname}}': cust['last_name'],
        '{{kunde_strasse}}': cust['street'],
        '{{kunde_plz}}': cust['zip'],
        '{{kunde_stadt}}': cust['city'],
        '{{kunde_nr}}': cust['customer_nr'],
        '{{kunde_email}}': cust['email'],
        '{{kunde_telefon}}': cust['phone'],
    }


def build_document(doc_type, doc_id):
    """Erzeugt das Word-Dokument aus der Vorlage und speichert es im Output-Ordner.
    Gibt (output_path, output_name) zurück; wirft DocGenError bei Problemen.
    Wird sowohl vom Download ("Word generieren") als auch vom E-Mail-Versand genutzt,
    damit beim Senden bei Bedarf automatisch (neu) generiert wird."""
    try:
        from docx import Document
    except ImportError:
        raise DocGenError('python-docx ist nicht installiert.')

    template_map = {
        'invoice': 'vorlage_rechnung.docx',
        'offer': 'vorlage_angebot.docx',
        'credit': 'vorlage_gutschrift.docx',
        'reminder': 'vorlage_mahnung.docx',
    }

    template_file = os.path.join(UPLOAD_FOLDER, template_map.get(doc_type, ''))
    if not os.path.exists(template_file):
        raise DocGenError(
            f'Vorlage "{template_map.get(doc_type)}" nicht gefunden im Ordner "vorlagen/".')

    db = get_db()
    s = get_settings()

    # Build replacement dict
    replacements = company_placeholders(s)

    output_name = ''

    if doc_type == 'invoice':
        inv = db.execute("SELECT * FROM invoices WHERE id=?", [doc_id]).fetchone()
        cust = db.execute("SELECT * FROM customers WHERE id=?", [inv['customer_id']]).fetchone()
        items = db.execute("SELECT * FROM invoice_items WHERE invoice_id=? ORDER BY position", [doc_id]).fetchall()
        replacements.update({
            '{{rechnung_nr}}': inv['invoice_nr'],
            '{{rechnung_datum}}': fmt_date(inv['date']),
            '{{faellig_datum}}': fmt_date(inv['due_date']),
            **customer_placeholders(cust),
            '{{gesamtbetrag}}': f"{inv['total']:.2f} €",
            '{{notizen}}': inv['notes'] or '',
        })
        output_name = f"{inv['invoice_nr']}.docx"

    elif doc_type == 'offer':
        off = db.execute("SELECT * FROM offers WHERE id=?", [doc_id]).fetchone()
        cust = db.execute("SELECT * FROM customers WHERE id=?", [off['customer_id']]).fetchone()
        items = db.execute("SELECT * FROM offer_items WHERE offer_id=? ORDER BY position", [doc_id]).fetchall()
        replacements.update({
            '{{angebot_nr}}': off['offer_nr'],
            '{{angebot_datum}}': fmt_date(off['date']),
            '{{gueltig_bis}}': fmt_date(off['valid_until']),
            **customer_placeholders(cust),
            '{{gesamtbetrag}}': f"{off['total']:.2f} €",
            '{{notizen}}': off['notes'] or '',
        })
        output_name = f"{off['offer_nr']}.docx"

    elif doc_type == 'credit':
        cr = db.execute("SELECT * FROM credits WHERE id=?", [doc_id]).fetchone()
        cust = db.execute("SELECT * FROM customers WHERE id=?", [cr['customer_id']]).fetchone()
        items = db.execute("SELECT * FROM credit_items WHERE credit_id=? ORDER BY position", [doc_id]).fetchall()
        replacements.update({
            '{{gutschrift_nr}}': cr['credit_nr'],
            '{{gutschrift_datum}}': fmt_date(cr['date']),
            **customer_placeholders(cust),
            '{{gesamtbetrag}}': f"{cr['total']:.2f} €",
            '{{notizen}}': cr['notes'] or '',
        })
        output_name = f"{cr['credit_nr']}.docx"

    elif doc_type == 'reminder':
        rem = db.execute('''SELECT r.*, i.invoice_nr, i.total as inv_total FROM reminders r
                            JOIN invoices i ON r.invoice_id=i.id WHERE r.id=?''', [doc_id]).fetchone()
        inv = db.execute("SELECT * FROM invoices WHERE id=?", [rem['invoice_id']]).fetchone()
        cust = db.execute("SELECT * FROM customers WHERE id=?", [inv['customer_id']]).fetchone()
        # Positionen der gemahnten Rechnung laden, damit {{positionen}} auch in der Mahnung funktioniert
        items = db.execute("SELECT * FROM invoice_items WHERE invoice_id=? ORDER BY position", [rem['invoice_id']]).fetchall()
        replacements.update({
            '{{mahnung_stufe}}': str(rem['level']),
            '{{mahnung_datum}}': fmt_date(rem['date']),
            '{{mahnung_frist}}': fmt_date(rem['due_date']),
            '{{mahngebuehr}}': f"{rem['fee']:.2f} €",
            '{{rechnung_nr}}': inv['invoice_nr'],
            '{{rechnung_datum}}': fmt_date(inv['date']),
            '{{rechnung_faellig_datum}}': fmt_date(inv['due_date']),
            '{{rechnung_betrag}}': f"{inv['total']:.2f} €",
            '{{gesamtbetrag}}': f"{inv['total'] + rem['fee']:.2f} €",
            **customer_placeholders(cust),
            '{{notizen}}': rem['notes'] or '',
        })
        output_name = f"Mahnung_{rem['level']}_{inv['invoice_nr']}.docx"

    db.close()

    # Process Word template
    doc = Document(template_file)

    # Platzhalter robust ersetzen (Body, Tabellen, Kopf-/Fußzeilen, run-übergreifend)
    replace_placeholders(doc, replacements)

    # Logo einfügen, wo {{logo}}-Marker steht (helle Variante für weißen Hintergrund);
    # ohne hinterlegtes Logo Marker entfernen, damit kein {{logo}}-Text stehen bleibt
    logo_path = get_logo_path(s, prefer='light')
    if logo_path:
        insert_logo(doc, logo_path, get_logo_width_cm(s))
    else:
        replace_placeholders(doc, {LOGO_PLACEHOLDER: ''})

    # Positionen einfügen ({{positionen}}-Marker -> Zeilen, Markerzeile wird entfernt)
    fill_positions(doc, items)

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    output_path = os.path.join(OUTPUT_FOLDER, output_name)
    doc.save(output_path)
    return output_path, output_name


@app.route('/generate/<doc_type>/<int:doc_id>')
@login_required
def generate_doc(doc_type, doc_id):
    try:
        output_path, output_name = build_document(doc_type, doc_id)
    except DocGenError as e:
        flash(str(e), 'error')
        return redirect(request.referrer or url_for('dashboard'))
    return send_file(output_path, as_attachment=True, download_name=output_name)


# Feste Fundorte für LibreOffice. Nötig, weil die systemd-Unit den PATH auf
# "/opt/faktura/venv/bin" setzt – dort steht soffice nicht, und shutil.which()
# sucht ausschließlich im PATH. Ohne diese Liste scheitert die PDF-Umwandlung im
# Dienst mit "nicht gefunden", obwohl LibreOffice installiert ist (auf der
# SSH-Shell mit vollem PATH funktioniert sie dagegen).
SOFFICE_CANDIDATES = (
    '/usr/bin/soffice',
    '/usr/bin/libreoffice',
    '/usr/lib/libreoffice/program/soffice',
    '/opt/libreoffice/program/soffice',
    '/snap/bin/libreoffice',
)


# Übliche Systempfade für den soffice-Aufruf. /usr/bin/soffice ist ein
# Shell-Skript und ruft dirname, basename, ls und sed auf – mit dem PATH der
# systemd-Unit (nur venv/bin) findet es diese Werkzeuge nicht und bricht mit
# "dirname: not found" ab.
SYSTEM_PATH_DIRS = ('/usr/local/bin', '/usr/bin', '/bin', '/usr/local/sbin',
                    '/usr/sbin', '/sbin')


def soffice_env():
    """Umgebung für den LibreOffice-Aufruf: PATH um die Systempfade ergänzt."""
    env = dict(os.environ)
    dirs = [d for d in env.get('PATH', '').split(os.pathsep) if d]
    dirs += [d for d in SYSTEM_PATH_DIRS if d not in dirs]
    env['PATH'] = os.pathsep.join(dirs)
    return env


def find_soffice():
    """Pfad zur LibreOffice-Binary oder None."""
    found = shutil.which('soffice') or shutil.which('libreoffice')
    if found:
        return found
    for path in SOFFICE_CANDIDATES:
        if os.path.isfile(path) and os.access(path, os.X_OK):
            return path
    return None


def convert_to_pdf(docx_path):
    """Wandelt eine .docx-Datei per LibreOffice (headless) in PDF um und gibt den
    PDF-Pfad zurück. Wirft DocGenError bei Fehlern (dann wird nichts versendet).
    www-data hat kein nutzbares HOME, deshalb ein eigenes, beschreibbares
    LibreOffice-Profil je Aufruf – das vermeidet auch das Single-Instance-Lock."""
    soffice = find_soffice()
    if not soffice:
        raise DocGenError('PDF-Umwandlung nicht möglich: LibreOffice (soffice) '
                          'wurde auf dem Server nicht gefunden.')
    out_dir = os.path.dirname(docx_path)
    profile = tempfile.mkdtemp(prefix='lo_profile_')
    pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
    # Eine ältere PDF gleichen Namens zuerst wegräumen. Sonst gilt sie unten als
    # Ergebnis, obwohl die Umwandlung gescheitert ist – und es würde ein
    # veralteter Beleg angehängt.
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    try:
        try:
            result = subprocess.run(
                [soffice, '--headless', '--nologo', '--norestore',
                 f'-env:UserInstallation=file://{profile}',
                 '--convert-to', 'pdf', '--outdir', out_dir, docx_path],
                capture_output=True, timeout=120, env=soffice_env())
        except subprocess.TimeoutExpired:
            raise DocGenError('PDF-Umwandlung hat zu lange gedauert.')
        if not os.path.exists(pdf_path):
            detail = result.stderr.decode('utf-8', 'replace').strip()[:200]
            raise DocGenError(f'PDF-Umwandlung fehlgeschlagen. {detail}'.strip())
        return pdf_path
    finally:
        shutil.rmtree(profile, ignore_errors=True)


def build_pdf(doc_type, doc_id):
    """Erzeugt den Beleg und wandelt ihn in PDF um.
    Gibt (Pfad, Dateiname) des PDFs zurück."""
    docx_path, docx_name = build_document(doc_type, doc_id)
    pdf_path = convert_to_pdf(docx_path)
    return pdf_path, os.path.splitext(docx_name)[0] + '.pdf'


@app.route('/generate-pdf/<doc_type>/<int:doc_id>')
@login_required
def generate_pdf(doc_type, doc_id):
    """Beleg als PDF herunterladen (fertiges Dokument zum Weitergeben).
    Zum Bearbeiten liefert /generate/... weiterhin die .docx."""
    try:
        pdf_path, pdf_name = build_pdf(doc_type, doc_id)
    except DocGenError as e:
        flash(str(e), 'error')
        return redirect(request.referrer or url_for('dashboard'))
    return send_file(pdf_path, as_attachment=True, download_name=pdf_name)


# ── E-Mail-Versand ──
#
# Zwei Verfahren, umschaltbar in den Einstellungen (settings.mail_method):
#   'smtp'  – klassischer SMTP-Versand (Benutzername/Passwort)
#   'graph' – Microsoft Graph, App-Registrierung im Entra Admin Center mit
#             der Anwendungsberechtigung Mail.Send (Client-Credentials-Flow)

GRAPH_TOKEN_URL = 'https://login.microsoftonline.com/{tenant}/oauth2/v2.0/token'
GRAPH_SCOPE = 'https://graph.microsoft.com/.default'
GRAPH_SENDMAIL_URL = 'https://graph.microsoft.com/v1.0/users/{sender}/sendMail'
# Entwurf im Postfach anlegen (statt sofort senden). Braucht die
# Anwendungsberechtigung Mail.ReadWrite – Mail.Send allein genügt dafür NICHT.
GRAPH_MESSAGES_URL = 'https://graph.microsoft.com/v1.0/users/{sender}/messages'
# sendMail überträgt Anhänge inline (base64). Größere Dateien bräuchten eine
# Upload-Session – für Rechnungen im docx-Format reicht das Limit deutlich.
GRAPH_MAX_ATTACHMENT_BYTES = 3 * 1024 * 1024

# Access-Token je (Tenant, Client-ID) zwischenspeichern; gilt rund 1 Stunde.
_graph_token_cache = {}


def setting(s, key, default=''):
    """Liest ein Feld aus den Einstellungen, auch wenn die Spalte fehlt."""
    try:
        value = s[key]
    except (KeyError, IndexError, TypeError):
        return default
    return default if value is None else value


def get_mail_method(s):
    return 'graph' if setting(s, 'mail_method', 'smtp') == 'graph' else 'smtp'


def get_graph_sender(s):
    """Postfach, aus dem gesendet wird (Fallback: SMTP-Absender, dann Firmen-E-Mail)."""
    return (setting(s, 'graph_sender') or setting(s, 'smtp_from')
            or setting(s, 'email')).strip()


def mail_is_configured(s):
    if get_mail_method(s) == 'graph':
        return all([setting(s, 'graph_tenant_id').strip(),
                    setting(s, 'graph_client_id').strip(),
                    setting(s, 'graph_client_secret'),
                    get_graph_sender(s)])
    return bool(setting(s, 'smtp_host'))


def mail_config_hint(s):
    if get_mail_method(s) == 'graph':
        return ('Microsoft-Graph-Versand ist nicht vollständig konfiguriert '
                '(Verzeichnis-ID, Anwendungs-ID, Client-Secret, Absender-Postfach).')
    return 'SMTP nicht konfiguriert. Bitte in Einstellungen hinterlegen.'


def _graph_error(response):
    """Fehlermeldung aus einer Graph-/Entra-Antwort lesbar aufbereiten."""
    try:
        data = response.json()
    except ValueError:
        return (response.text or '').strip()[:300]
    if isinstance(data.get('error'), dict):
        return data['error'].get('message', str(data['error']))[:300]
    if data.get('error_description'):
        return str(data['error_description']).splitlines()[0][:300]
    return str(data)[:300]


def get_graph_token(s):
    """Access-Token per Client-Credentials-Flow holen (mit Cache)."""
    tenant = setting(s, 'graph_tenant_id').strip()
    client_id = setting(s, 'graph_client_id').strip()
    secret = setting(s, 'graph_client_secret')

    cache_key = (tenant, client_id)
    cached = _graph_token_cache.get(cache_key)
    if cached and cached['expires_at'] > time.time() + 60:
        return cached['token']

    response = requests.post(
        GRAPH_TOKEN_URL.format(tenant=tenant),
        data={'client_id': client_id, 'client_secret': secret,
              'scope': GRAPH_SCOPE, 'grant_type': 'client_credentials'},
        timeout=30)
    if response.status_code != 200:
        raise RuntimeError(f"Token-Abruf fehlgeschlagen ({response.status_code}): "
                           f"{_graph_error(response)}")

    data = response.json()
    _graph_token_cache[cache_key] = {
        'token': data['access_token'],
        'expires_at': time.time() + int(data.get('expires_in', 3600)),
    }
    return data['access_token']


def attachment_mimetype(filename):
    return 'application/pdf' if filename.lower().endswith('.pdf') else 'application/octet-stream'


def build_graph_message(recipient, subject, body, attachments):
    """Baut das Graph-Nachrichtenobjekt (für sendMail und für Entwürfe)."""
    message = {
        'subject': subject,
        'body': {'contentType': 'Text', 'content': body},
        'toRecipients': [{'emailAddress': {'address': recipient}}],
    }
    if attachments:
        message['attachments'] = []
        for filename, content in attachments:
            if len(content) > GRAPH_MAX_ATTACHMENT_BYTES:
                raise RuntimeError(
                    f"Anhang {filename} ist zu groß für den Graph-Versand "
                    f"(max. {GRAPH_MAX_ATTACHMENT_BYTES // (1024 * 1024)} MB).")
            message['attachments'].append({
                '@odata.type': '#microsoft.graph.fileAttachment',
                'name': filename,
                'contentType': attachment_mimetype(filename),
                'contentBytes': base64.b64encode(content).decode('ascii'),
            })
    return message


def graph_token_roles(s):
    """Liest die Anwendungsberechtigungen (roles) aus dem Access-Token.
    Gibt None zurück, wenn sie sich nicht ermitteln lassen."""
    try:
        payload = get_graph_token(s).split('.')[1]
        payload += '=' * (-len(payload) % 4)
        claims = json.loads(base64.urlsafe_b64decode(payload))
        roles = claims.get('roles')
        return roles if isinstance(roles, list) else None
    except Exception:
        return None


def graph_can_draft(s):
    """Darf die App-Registrierung Entwürfe im Postfach anlegen? Bei unbekannten
    Rollen wird es versucht – ein Fehlschlag führt zum .eml-Download."""
    roles = graph_token_roles(s)
    return roles is None or 'Mail.ReadWrite' in roles


def create_graph_draft(s, recipient, subject, body, attachments):
    """Legt die E-Mail als Entwurf im Postfach ab
    (POST /users/{sender}/messages). Es wird nichts versendet."""
    response = requests.post(
        GRAPH_MESSAGES_URL.format(sender=get_graph_sender(s)),
        headers={'Authorization': f'Bearer {get_graph_token(s)}',
                 'Content-Type': 'application/json'},
        json=build_graph_message(recipient, subject, body, attachments), timeout=60)
    if response.status_code not in (200, 201):
        raise RuntimeError(f"Entwurf konnte nicht angelegt werden "
                           f"({response.status_code}): {_graph_error(response)}")


def build_eml(s, recipient, subject, body, attachments):
    """Baut die E-Mail als .eml-Datei (Bytes). Der Header X-Unsent sorgt dafür,
    dass Outlook sie als noch nicht gesendeten Entwurf mit Senden-Button öffnet."""
    msg = MIMEMultipart()
    msg['From'] = (get_graph_sender(s) if get_mail_method(s) == 'graph'
                   else setting(s, 'smtp_from'))
    msg['To'] = recipient
    msg['Subject'] = subject
    msg['X-Unsent'] = '1'
    msg.attach(MIMEText(body, 'plain', 'utf-8'))

    for filename, content in attachments:
        subtype = 'pdf' if filename.lower().endswith('.pdf') else 'octet-stream'
        part = MIMEBase('application', subtype)
        part.set_payload(content)
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
        msg.attach(part)

    return msg.as_bytes()


def send_mail_graph(s, recipient, subject, body, attachments):
    """Versand über Microsoft Graph (POST /users/{sender}/sendMail)."""
    sender = get_graph_sender(s)
    payload = {'message': build_graph_message(recipient, subject, body, attachments),
               'saveToSentItems': bool(setting(s, 'graph_save_sent', 1))}
    response = requests.post(
        GRAPH_SENDMAIL_URL.format(sender=sender),
        headers={'Authorization': f'Bearer {get_graph_token(s)}',
                 'Content-Type': 'application/json'},
        json=payload, timeout=60)
    # Erfolg ist 202 Accepted ohne Inhalt
    if response.status_code not in (200, 202):
        raise RuntimeError(f"Graph-Versand fehlgeschlagen ({response.status_code}): "
                           f"{_graph_error(response)}")


def send_mail_smtp(s, recipient, subject, body, attachments):
    """Versand über SMTP (STARTTLS)."""
    msg = MIMEMultipart()
    msg['From'] = setting(s, 'smtp_from')
    msg['To'] = recipient
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain', 'utf-8'))

    for filename, content in attachments:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(content)
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
        msg.attach(part)

    with smtplib.SMTP(setting(s, 'smtp_host'), int(setting(s, 'smtp_port', 587) or 587)) as server:
        server.starttls()
        server.login(setting(s, 'smtp_user'), setting(s, 'smtp_pass'))
        server.send_message(msg)


def send_mail(s, recipient, subject, body, attachments=()):
    """Verschickt eine E-Mail mit dem in den Einstellungen gewählten Verfahren."""
    attachments = list(attachments)
    if get_mail_method(s) == 'graph':
        send_mail_graph(s, recipient, subject, body, attachments)
    else:
        send_mail_smtp(s, recipient, subject, body, attachments)


def mail_anrede(cust):
    """Passende Briefanrede aus den Kundendaten."""
    salutation = (cust['salutation'] or '').strip()
    last_name = (cust['last_name'] or '').strip()
    if salutation == 'Herr' and last_name:
        return f"Sehr geehrter Herr {last_name},"
    if salutation == 'Frau' and last_name:
        return f"Sehr geehrte Frau {last_name},"
    if last_name:
        # Kein Geschlecht hinterlegt – neutrale, aber persönliche Anrede
        vorname = (cust['first_name'] or '').strip()
        return f"Guten Tag {(vorname + ' ' + last_name).strip()},"
    return "Sehr geehrte Damen und Herren,"


def mail_signatur(s):
    """Grußformel mit Absenderdaten – ohne doppelte Zeilen."""
    zeilen = ['Mit freundlichen Grüßen', '']
    for wert in (setting(s, 'owner_name').strip(), setting(s, 'company_name').strip()):
        if wert and wert not in zeilen:
            zeilen.append(wert)
    for wert in (setting(s, 'phone').strip(), setting(s, 'email').strip()):
        if wert:
            zeilen.append(wert)
    return '\n'.join(zeilen)


def apply_placeholders(text, ph):
    """Ersetzt {{platzhalter}} im Text durch die Werte aus dem Wörterbuch ph."""
    for key, val in ph.items():
        text = text.replace(key, '' if val is None else str(val))
    return text


def get_email_template(db, doc_type):
    """Lädt die E-Mail-Vorlage (subject, body) für die Belegart; fällt auf die
    Standardvorlage zurück, falls (noch) keine gespeichert ist."""
    row = db.execute("SELECT subject, body FROM email_templates WHERE doc_type=?",
                     [doc_type]).fetchone()
    if row:
        return row['subject'], row['body']
    default = EMAIL_TEMPLATE_DEFAULTS.get(doc_type, {'subject': '', 'body': ''})
    return default['subject'], default['body']


@app.route('/settings/test-mail', methods=['POST'])
@login_required
def test_mail():
    """Test-E-Mail mit den gespeicherten Einstellungen verschicken."""
    s = get_settings()
    if not mail_is_configured(s):
        flash(mail_config_hint(s), 'error')
        return redirect(url_for('settings'))

    recipient = (request.form.get('test_recipient') or '').strip()
    if not recipient:
        recipient = (setting(s, 'graph_sender') if get_mail_method(s) == 'graph'
                     else setting(s, 'smtp_from')) or setting(s, 'email')
    if not recipient:
        flash('Keine Empfängeradresse für den Test angegeben.', 'error')
        return redirect(url_for('settings'))

    verfahren = 'Microsoft Graph' if get_mail_method(s) == 'graph' else 'SMTP'
    try:
        send_mail(s, recipient, f'Testmail von {setting(s, "company_name") or "Micro-Fakt"}',
                  f'Diese Test-E-Mail wurde über {verfahren} verschickt.\n'
                  'Der E-Mail-Versand ist damit korrekt eingerichtet.')
        flash(f'Test-E-Mail über {verfahren} an {recipient} gesendet.', 'success')
    except Exception as e:
        flash(f'Test-E-Mail fehlgeschlagen: {e}', 'error')
    return redirect(url_for('settings'))


@app.route('/draft-email/<doc_type>/<int:doc_id>', methods=['POST'])
@login_required
def draft_email(doc_type, doc_id):
    """Bereitet die E-Mail zum Beleg vor – Empfänger, Betreff, Anschreiben und
    das PDF als Anhang – und verschickt sie BEWUSST NICHT. Geprüft und gesendet
    wird von Hand:
      * Microsoft Graph mit Mail.ReadWrite: Entwurf landet im Postfach
      * sonst: .eml-Datei zum Download, öffnet in Outlook als Entwurf
    Der Belegstatus bleibt unverändert, weil noch nichts versendet wurde."""
    s = get_settings()
    if not mail_is_configured(s):
        flash(mail_config_hint(s), 'error')
        return redirect(request.referrer or url_for('dashboard'))

    db = get_db()

    # Kunde und belegspezifische Platzhalter je Dokumenttyp
    if doc_type == 'invoice':
        doc = db.execute("SELECT * FROM invoices WHERE id=?", [doc_id]).fetchone()
        cust = db.execute("SELECT * FROM customers WHERE id=?", [doc['customer_id']]).fetchone()
        ph = {
            '{{rechnung_nr}}': doc['invoice_nr'],
            '{{rechnung_datum}}': fmt_date(doc['date']),
            '{{faellig_datum}}': fmt_date(doc['due_date']),
            '{{betrag}}': f"{doc['total']:.2f} €",
        }
    elif doc_type == 'offer':
        doc = db.execute("SELECT * FROM offers WHERE id=?", [doc_id]).fetchone()
        cust = db.execute("SELECT * FROM customers WHERE id=?", [doc['customer_id']]).fetchone()
        ph = {
            '{{angebot_nr}}': doc['offer_nr'],
            '{{angebot_datum}}': fmt_date(doc['date']),
            '{{gueltig_bis}}': fmt_date(doc['valid_until']),
            '{{betrag}}': f"{doc['total']:.2f} €",
        }
    elif doc_type == 'credit':
        doc = db.execute("SELECT * FROM credits WHERE id=?", [doc_id]).fetchone()
        cust = db.execute("SELECT * FROM customers WHERE id=?", [doc['customer_id']]).fetchone()
        ph = {
            '{{gutschrift_nr}}': doc['credit_nr'],
            '{{gutschrift_datum}}': fmt_date(doc['date']),
            '{{betrag}}': f"{doc['total']:.2f} €",
        }
    elif doc_type == 'reminder':
        rem = db.execute("SELECT * FROM reminders WHERE id=?", [doc_id]).fetchone()
        inv = db.execute("SELECT * FROM invoices WHERE id=?", [rem['invoice_id']]).fetchone()
        cust = db.execute("SELECT * FROM customers WHERE id=?", [inv['customer_id']]).fetchone()
        ph = {
            '{{mahnung_stufe}}': str(rem['level']),
            '{{mahnung_frist}}': fmt_date(rem['due_date']),
            '{{mahngebuehr}}': f"{rem['fee']:.2f} €",
            '{{rechnung_nr}}': inv['invoice_nr'],
            '{{rechnung_datum}}': fmt_date(inv['date']),
            '{{rechnung_faellig_datum}}': fmt_date(inv['due_date']),
            '{{betrag}}': f"{inv['total'] + rem['fee']:.2f} €",
        }
    else:
        flash('Unbekannter Dokumenttyp.', 'error')
        return redirect(url_for('dashboard'))

    recipient = cust['email']
    if not recipient:
        flash('Kunde hat keine E-Mail-Adresse hinterlegt.', 'error')
        db.close()
        return redirect(request.referrer or url_for('dashboard'))

    # Gemeinsame Platzhalter: dieselben Einzelfelder wie in den Word-Vorlagen,
    # dazu die beiden zusammengesetzten {{anrede}} und {{signatur}}.
    ph.update(company_placeholders(s))
    ph.update(customer_placeholders(cust))
    ph.update({
        '{{anrede}}': mail_anrede(cust),
        '{{signatur}}': mail_signatur(s),
    })

    subj_tmpl, body_tmpl = get_email_template(db, doc_type)
    subject = apply_placeholders(subj_tmpl, ph)
    body = apply_placeholders(body_tmpl, ph)

    # Dokument bei Bedarf automatisch erzeugen (Output-Ordner ist nur temporär),
    # damit der Entwurf nicht voraussetzt, dass vorher "Word generieren" geklickt
    # wurde. Angehängt wird ausschließlich das PDF, nie die .docx.
    try:
        pdf_path, pdf_name = build_pdf(doc_type, doc_id)
    except DocGenError as e:
        flash(str(e), 'error')
        db.close()
        return redirect(request.referrer or url_for('dashboard'))
    db.close()

    with open(pdf_path, 'rb') as f:
        attachments = [(pdf_name, f.read())]

    # Bevorzugt als Entwurf ins Postfach – dort kann die Mail vor dem Senden
    # geprüft werden. Fehlt die Berechtigung dafür, gibt es die .eml-Datei.
    if get_mail_method(s) == 'graph' and graph_can_draft(s):
        try:
            create_graph_draft(s, recipient, subject, body, attachments)
            # Bewusst nur eine Meldung, kein Öffnen: der Entwurf wird im
            # Desktop-Outlook bearbeitet. Ein Link auf Outlook im Web führt dort
            # ins Postfach des im Browser angemeldeten Kontos – also womöglich
            # ins falsche.
            flash('E-Mail im Outlook-Ordner "Entwürfe" erstellt.', 'success')
            return redirect(request.referrer or url_for('dashboard'))
        except Exception as e:
            # Kein Entwurf möglich (z.B. fehlende Berechtigung) – .eml ausliefern
            app.logger.warning('Graph-Entwurf fehlgeschlagen, .eml-Fallback: %s', e)

    # Kein Flash beim Download: die Meldung würde erst auf der nächsten Seite
    # erscheinen. Die heruntergeladene Datei ist die Rückmeldung.
    eml_name = os.path.splitext(pdf_name)[0] + '.eml'
    return send_file(io.BytesIO(build_eml(s, recipient, subject, body, attachments)),
                     as_attachment=True, download_name=eml_name,
                     mimetype='message/rfc822')


# ── API for article data ──

@app.route('/api/article/<int:id>')
@login_required
def api_article(id):
    db = get_db()
    a = db.execute("SELECT * FROM articles WHERE id=?", [id]).fetchone()
    db.close()
    if a:
        return jsonify({'name': a['name'], 'description': a['description'],
                        'unit': a['unit'], 'price': a['price']})
    return jsonify({}), 404


# ── Init & Run ──


@app.route('/vorlagen/muster/<doc_type>')
@login_required
def vorlage_muster(doc_type):
    try:
        from docx import Document as DocxDocument
    except ImportError:
        flash('python-docx ist nicht installiert.', 'error')
        return redirect(url_for('vorlagen'))

    template_map = {
        'rechnung': 'vorlage_rechnung.docx',
        'angebot': 'vorlage_angebot.docx',
        'gutschrift': 'vorlage_gutschrift.docx',
        'mahnung': 'vorlage_mahnung.docx',
    }
    if doc_type not in template_map:
        flash('Unbekannter Dokumenttyp.', 'error')
        return redirect(url_for('vorlagen'))

    template_file = os.path.join(UPLOAD_FOLDER, template_map[doc_type])
    if not os.path.exists(template_file):
        flash(f'Vorlage "{template_map[doc_type]}" nicht gefunden.', 'error')
        return redirect(url_for('vorlagen'))

    # Beispieldaten
    beispiel = {
        '{{firma}}': 'Dirk Hildebrand - App Entwicklung',
        '{{inhaber}}': 'Dirk Hildebrand',
        '{{firma_strasse}}': 'Hirschbergstr. 4',
        '{{firma_plz}}': '34123',
        '{{firma_stadt}}': 'Kassel',
        '{{firma_telefon}}': '+49 1512 8225666',
        '{{firma_email}}': 'dirk@dirkhildebrand.de',
        '{{steuernummer}}': '026 123 45678',
        '{{bank}}': 'Stadtsparkasse Grebenstein',
        '{{iban}}': 'DE89 3704 0044 0532 0130 00',
        '{{bic}}': 'COBADEFFXXX',
        '{{kleinunternehmer}}': 'Gemäß §19 UStG wird keine Umsatzsteuer berechnet.',
        '{{kunde_firma}}': 'Mustermann GmbH',
        '{{kunde_anrede}}': 'Herr',
        '{{kunde_vorname}}': 'Max',
        '{{kunde_nachname}}': 'Mustermann',
        '{{kunde_strasse}}': 'Beispielweg 42',
        '{{kunde_plz}}': '34117',
        '{{kunde_stadt}}': 'Kassel',
        '{{kunde_nr}}': 'K-0001',
        '{{kunde_email}}': 'max.mustermann@example.com',
        '{{kunde_telefon}}': '+49 561 1234567',
        '{{notizen}}': '',
    }

    if doc_type == 'rechnung':
        beispiel.update({
            '{{rechnung_nr}}': 'RE-2604001',
            '{{rechnung_datum}}': '10.04.2026',
            '{{faellig_datum}}': '24.04.2026',
            '{{gesamtbetrag}}': '2.850,00 €',
        })
        output_name = 'Muster_Rechnung.docx'
    elif doc_type == 'angebot':
        beispiel.update({
            '{{angebot_nr}}': 'AN-2604001',
            '{{angebot_datum}}': '10.04.2026',
            '{{gueltig_bis}}': '10.05.2026',
            '{{gesamtbetrag}}': '2.850,00 €',
        })
        output_name = 'Muster_Angebot.docx'
    elif doc_type == 'gutschrift':
        beispiel.update({
            '{{gutschrift_nr}}': 'GU-2604001',
            '{{gutschrift_datum}}': '10.04.2026',
            '{{gesamtbetrag}}': '500,00 €',
        })
        output_name = 'Muster_Gutschrift.docx'
    elif doc_type == 'mahnung':
        beispiel.update({
            '{{mahnung_stufe}}': '1',
            '{{mahnung_datum}}': '10.04.2026',
            '{{mahnung_frist}}': '17.04.2026',
            '{{mahngebuehr}}': '5,00 €',
            '{{rechnung_nr}}': 'RE-2603001',
            '{{rechnung_datum}}': '01.03.2026',
            '{{rechnung_faellig_datum}}': '15.03.2026',
            '{{rechnung_betrag}}': '1.200,00 €',
            '{{gesamtbetrag}}': '1.205,00 €',
        })
        output_name = 'Muster_Mahnung.docx'

    doc = DocxDocument(template_file)

    # Platzhalter robust ersetzen (Body, Tabellen, Kopf-/Fußzeilen, run-übergreifend)
    replace_placeholders(doc, beispiel)

    # Logo einfügen, wo {{logo}}-Marker steht (helle Variante für weißen Hintergrund);
    # ohne hinterlegtes Logo Marker entfernen, damit kein {{logo}}-Text stehen bleibt
    muster_settings = get_settings()
    logo_path = get_logo_path(muster_settings, prefer='light')
    if logo_path:
        insert_logo(doc, logo_path, get_logo_width_cm(muster_settings))
    else:
        replace_placeholders(doc, {LOGO_PLACEHOLDER: ''})

    # Positionen einfügen (falls die Vorlage einen {{positionen}}-Marker hat)
    muster_positionen = [
        {'position': 1, 'description': 'Power Apps Entwicklung', 'quantity': 10, 'unit': 'Stunde(n)', 'price': 85.00, 'total': 850.00},
        {'position': 2, 'description': 'Datenbank-Design und Einrichtung', 'quantity': 1, 'unit': 'Pauschale', 'price': 500.00, 'total': 500.00},
        {'position': 3, 'description': 'Monatliche Wartung & Support', 'quantity': 3, 'unit': 'pro Monat', 'price': 500.00, 'total': 1500.00},
    ]
    fill_positions(doc, muster_positionen)

    output_path = os.path.join(OUTPUT_FOLDER, output_name)
    doc.save(output_path)
    return send_file(output_path, as_attachment=True, download_name=output_name)

# ── Vorlagen ──

@app.route('/vorlagen')
@login_required
def vorlagen():
    templates = [
        {'label': 'Rechnung', 'filename': 'vorlage_rechnung.docx', 'type': 'rechnung'},
        {'label': 'Angebot', 'filename': 'vorlage_angebot.docx', 'type': 'angebot'},
        {'label': 'Gutschrift', 'filename': 'vorlage_gutschrift.docx', 'type': 'gutschrift'},
        {'label': 'Mahnung', 'filename': 'vorlage_mahnung.docx', 'type': 'mahnung'},
    ]
    for t in templates:
        t['exists'] = os.path.exists(os.path.join(UPLOAD_FOLDER, t['filename']))
    return render_template('vorlagen.html', vorlagen=templates)


# ── E-Mail-Vorlagen ──

# Belegarten mit Anzeigename und den in der jeweiligen E-Mail verfügbaren Platzhaltern
EMAIL_TEMPLATE_META = [
    {'doc_type': 'invoice', 'label': 'Rechnung',
     'placeholders': ['{{rechnung_nr}}', '{{rechnung_datum}}',
                      '{{faellig_datum}}', '{{betrag}}']},
    {'doc_type': 'offer', 'label': 'Angebot',
     'placeholders': ['{{angebot_nr}}', '{{angebot_datum}}',
                      '{{gueltig_bis}}', '{{betrag}}']},
    {'doc_type': 'credit', 'label': 'Gutschrift',
     'placeholders': ['{{gutschrift_nr}}', '{{gutschrift_datum}}',
                      '{{betrag}}']},
    {'doc_type': 'reminder', 'label': 'Mahnung',
     'placeholders': ['{{mahnung_stufe}}', '{{mahnung_frist}}',
                      '{{mahngebuehr}}', '{{rechnung_nr}}', '{{rechnung_datum}}',
                      '{{rechnung_faellig_datum}}', '{{betrag}}']},
]

# Zusätzlich in jeder Vorlage nutzbar
# Einzelfelder für die E-Mail-Vorlagen – gleiche Namen wie in den Word-Vorlagen,
# damit sich Anschreiben und Dokument aus denselben Bausteinen zusammensetzen
# lassen. {{anrede}} und {{signatur}} sind die einzigen zusammengesetzten.
EMAIL_CUSTOMER_PLACEHOLDERS = ['{{kunde_anrede}}', '{{kunde_vorname}}', '{{kunde_nachname}}',
                               '{{kunde_firma}}', '{{kunde_nr}}', '{{kunde_strasse}}',
                               '{{kunde_plz}}', '{{kunde_stadt}}', '{{kunde_email}}',
                               '{{kunde_telefon}}']
EMAIL_COMPANY_PLACEHOLDERS = ['{{inhaber}}', '{{firma}}', '{{firma_strasse}}', '{{firma_plz}}',
                              '{{firma_stadt}}', '{{firma_telefon}}', '{{firma_email}}',
                              '{{steuernummer}}', '{{bank}}', '{{iban}}', '{{bic}}',
                              '{{kleinunternehmer}}']


@app.route('/email-vorlagen')
@login_required
def email_templates():
    db = get_db()
    rows = {r['doc_type']: r for r in db.execute("SELECT * FROM email_templates").fetchall()}
    db.close()
    items = []
    for meta in EMAIL_TEMPLATE_META:
        row = rows.get(meta['doc_type'])
        default = EMAIL_TEMPLATE_DEFAULTS.get(meta['doc_type'], {'subject': '', 'body': ''})
        items.append({
            'doc_type': meta['doc_type'],
            'label': meta['label'],
            'placeholder_groups': [
                {'label': 'Beleg', 'items': meta['placeholders']},
                {'label': 'Kunde', 'items': EMAIL_CUSTOMER_PLACEHOLDERS},
                {'label': 'Absender', 'items': EMAIL_COMPANY_PLACEHOLDERS},
                {'label': 'Fertige Bausteine', 'items': ['{{anrede}}', '{{signatur}}']},
            ],
            'subject': row['subject'] if row else default['subject'],
            'body': row['body'] if row else default['body'],
        })
    return render_template('email_templates.html', templates=items)


@app.route('/email-vorlagen/save/<doc_type>', methods=['POST'])
@login_required
def email_template_save(doc_type):
    if doc_type not in EMAIL_TEMPLATE_DEFAULTS:
        flash('Unbekannte Belegart.', 'error')
        return redirect(url_for('email_templates'))
    subject = request.form.get('subject', '').strip()
    body = request.form.get('body', '')
    db = get_db()
    db.execute("INSERT OR REPLACE INTO email_templates (doc_type, subject, body) VALUES (?, ?, ?)",
               [doc_type, subject, body])
    db.commit()
    db.close()
    flash('E-Mail-Vorlage gespeichert.', 'success')
    return redirect(url_for('email_templates'))


@app.route('/email-vorlagen/reset/<doc_type>', methods=['POST'])
@login_required
def email_template_reset(doc_type):
    default = EMAIL_TEMPLATE_DEFAULTS.get(doc_type)
    if not default:
        flash('Unbekannte Belegart.', 'error')
        return redirect(url_for('email_templates'))
    db = get_db()
    db.execute("INSERT OR REPLACE INTO email_templates (doc_type, subject, body) VALUES (?, ?, ?)",
               [doc_type, default['subject'], default['body']])
    db.commit()
    db.close()
    flash('E-Mail-Vorlage auf Standard zurückgesetzt.', 'success')
    return redirect(url_for('email_templates'))


@app.route('/vorlagen/download/<filename>')
@login_required
def vorlage_download(filename):
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True, download_name=filename)
    flash('Vorlage nicht gefunden.', 'error')
    return redirect(url_for('vorlagen'))


@app.route('/vorlagen/upload', methods=['POST'])
@login_required
def vorlage_upload():
    target = request.form.get('target', '')
    allowed = ['vorlage_rechnung.docx', 'vorlage_angebot.docx',
               'vorlage_gutschrift.docx', 'vorlage_mahnung.docx']
    if target not in allowed:
        flash('Ungültiger Dateiname.', 'error')
        return redirect(url_for('vorlagen'))
    file = request.files.get('file')
    if not file or not file.filename.endswith('.docx'):
        flash('Bitte eine .docx-Datei auswählen.', 'error')
        return redirect(url_for('vorlagen'))
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    try:
        file.save(os.path.join(UPLOAD_FOLDER, target))
    except OSError as e:
        # Typischer Fall: die Datei kam per git pull als root ins Verzeichnis und
        # gehört deshalb nicht dem Dienst-User (www-data).
        flash(f'Vorlage konnte nicht gespeichert werden: {e}. '
              f'Rechte prüfen: chown -R www-data:www-data {UPLOAD_FOLDER}', 'error')
        return redirect(url_for('vorlagen'))
    flash(f'Vorlage "{target}" hochgeladen.', 'success')
    return redirect(url_for('vorlagen'))

if __name__ == '__main__':
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    init_db()
    app.run(debug=True, port=5000)
