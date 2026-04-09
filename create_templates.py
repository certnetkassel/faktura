"""Erstellt Word-Vorlagen für Rechnung, Angebot, Gutschrift und Mahnung."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def style_doc(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)


def add_header_block(doc):
    # Firmenname
    p = doc.add_paragraph()
    run = p.add_run('{{firma}}')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x6C, 0x5C, 0xE7)
    run.bold = True
    p.space_after = Pt(2)

    # Firmendaten
    p = doc.add_paragraph()
    p.style.font.size = Pt(8)
    run = p.add_run('{{firma_strasse}} · {{firma_plz}} {{firma_stadt}} · Tel: {{firma_telefon}} · {{firma_email}}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.space_after = Pt(6)

    # Trennlinie
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    # Horizontal line via bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '6C5CE7',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_address_block(doc):
    p = doc.add_paragraph()
    run = p.add_run('{{kunde_anrede}} {{kunde_vorname}} {{kunde_nachname}}')
    run.font.size = Pt(11)
    run.bold = True
    if True:
        p2 = doc.add_paragraph()
        run2 = p2.add_run('{{kunde_firma}}')
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p2.space_after = Pt(0)

    p3 = doc.add_paragraph('{{kunde_strasse}}')
    p3.space_after = Pt(0)
    p4 = doc.add_paragraph('{{kunde_plz}} {{kunde_stadt}}')
    p4.space_after = Pt(20)


def add_positions_table(doc):
    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    headers = ['Pos.', 'Beschreibung', 'Menge', 'Einheit', 'Einzelpreis', 'Gesamt']
    widths = [Cm(1), Cm(7), Cm(2), Cm(2), Cm(2.5), Cm(2.5)]

    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.width = width
        set_cell_shading(cell, '6C5CE7')

    # Placeholder row
    marker_cell = table.rows[1].cells[0]
    marker_cell.text = '{{positionen}}'
    marker_cell.paragraphs[0].runs[0].font.size = Pt(9)
    for i in range(1, 6):
        table.rows[1].cells[i].text = ''

    doc.add_paragraph()  # spacing


def add_total_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Gesamtbetrag: {{gesamtbetrag}}')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x6C, 0x5C, 0xE7)
    p.space_after = Pt(6)


def add_kleinunternehmer(doc):
    p = doc.add_paragraph()
    run = p.add_run('{{kleinunternehmer}}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True
    p.space_after = Pt(12)


def add_bank_footer(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    top = pBdr.makeelement(qn('w:top'), {
        qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '4', qn('w:color'): 'CCCCCC',
    })
    pBdr.append(top)
    pPr.append(pBdr)
    p.space_before = Pt(20)

    run = p.add_run('Bankverbindung: {{bank}} · IBAN: {{iban}} · BIC: {{bic}}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p2 = doc.add_paragraph()
    run2 = p2.add_run('{{firma}} · {{inhaber}} · Steuernummer: {{steuernummer}}')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def create_invoice():
    doc = Document()
    style_doc(doc)
    add_header_block(doc)
    add_address_block(doc)

    # Title
    p = doc.add_paragraph()
    run = p.add_run('Rechnung {{rechnung_nr}}')
    run.font.size = Pt(16)
    run.bold = True
    p.space_after = Pt(4)

    # Meta
    meta = doc.add_paragraph()
    run = meta.add_run('Rechnungsdatum: {{rechnung_datum}}     Fällig bis: {{faellig_datum}}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    meta.space_after = Pt(16)

    # Intro
    p = doc.add_paragraph('Sehr geehrte(r) {{kunde_anrede}} {{kunde_nachname}},')
    p.space_after = Pt(4)
    p = doc.add_paragraph('für die erbrachten Leistungen erlaube ich mir, folgenden Betrag in Rechnung zu stellen:')
    p.space_after = Pt(12)

    add_positions_table(doc)
    add_total_block(doc)
    add_kleinunternehmer(doc)

    # Payment info
    p = doc.add_paragraph()
    run = p.add_run('Bitte überweisen Sie den Gesamtbetrag bis zum {{faellig_datum}} auf das unten genannte Konto.')
    run.font.size = Pt(10)
    p.space_after = Pt(4)

    # Notes
    p = doc.add_paragraph('{{notizen}}')
    p.space_after = Pt(8)

    p = doc.add_paragraph('Mit freundlichen Grüßen')
    p.space_after = Pt(4)
    p = doc.add_paragraph('{{inhaber}}')

    add_bank_footer(doc)

    doc.save(os.path.join('vorlagen', 'vorlage_rechnung.docx'))
    print('  vorlage_rechnung.docx erstellt')


def create_offer():
    doc = Document()
    style_doc(doc)
    add_header_block(doc)
    add_address_block(doc)

    p = doc.add_paragraph()
    run = p.add_run('Angebot {{angebot_nr}}')
    run.font.size = Pt(16)
    run.bold = True
    p.space_after = Pt(4)

    meta = doc.add_paragraph()
    run = meta.add_run('Datum: {{angebot_datum}}     Gültig bis: {{gueltig_bis}}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    meta.space_after = Pt(16)

    p = doc.add_paragraph('Sehr geehrte(r) {{kunde_anrede}} {{kunde_nachname}},')
    p.space_after = Pt(4)
    p = doc.add_paragraph('vielen Dank für Ihre Anfrage. Gerne unterbreite ich Ihnen folgendes Angebot:')
    p.space_after = Pt(12)

    add_positions_table(doc)
    add_total_block(doc)
    add_kleinunternehmer(doc)

    p = doc.add_paragraph('{{notizen}}')
    p.space_after = Pt(8)

    p = doc.add_paragraph('Dieses Angebot ist gültig bis zum {{gueltig_bis}}. Bei Fragen stehe ich Ihnen gerne zur Verfügung.')
    p.space_after = Pt(12)

    p = doc.add_paragraph('Mit freundlichen Grüßen')
    p.space_after = Pt(4)
    p = doc.add_paragraph('{{inhaber}}')

    add_bank_footer(doc)

    doc.save(os.path.join('vorlagen', 'vorlage_angebot.docx'))
    print('  vorlage_angebot.docx erstellt')


def create_credit():
    doc = Document()
    style_doc(doc)
    add_header_block(doc)
    add_address_block(doc)

    p = doc.add_paragraph()
    run = p.add_run('Gutschrift {{gutschrift_nr}}')
    run.font.size = Pt(16)
    run.bold = True
    p.space_after = Pt(4)

    meta = doc.add_paragraph()
    run = meta.add_run('Datum: {{gutschrift_datum}}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    meta.space_after = Pt(16)

    p = doc.add_paragraph('Sehr geehrte(r) {{kunde_anrede}} {{kunde_nachname}},')
    p.space_after = Pt(4)
    p = doc.add_paragraph('hiermit erhalten Sie eine Gutschrift über folgende Positionen:')
    p.space_after = Pt(12)

    add_positions_table(doc)
    add_total_block(doc)
    add_kleinunternehmer(doc)

    p = doc.add_paragraph('Der Betrag wird Ihnen auf das uns bekannte Konto überwiesen.')
    p.space_after = Pt(8)

    p = doc.add_paragraph('{{notizen}}')
    p.space_after = Pt(8)

    p = doc.add_paragraph('Mit freundlichen Grüßen')
    p.space_after = Pt(4)
    p = doc.add_paragraph('{{inhaber}}')

    add_bank_footer(doc)

    doc.save(os.path.join('vorlagen', 'vorlage_gutschrift.docx'))
    print('  vorlage_gutschrift.docx erstellt')


def create_reminder():
    doc = Document()
    style_doc(doc)
    add_header_block(doc)
    add_address_block(doc)

    p = doc.add_paragraph()
    run = p.add_run('{{mahnung_stufe}}. Mahnung')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0xE1, 0x70, 0x55)
    p.space_after = Pt(4)

    meta = doc.add_paragraph()
    run = meta.add_run('Datum: {{mahnung_datum}}     Bezug: Rechnung {{rechnung_nr}} vom {{rechnung_datum}}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    meta.space_after = Pt(16)

    p = doc.add_paragraph('Sehr geehrte(r) {{kunde_anrede}} {{kunde_nachname}},')
    p.space_after = Pt(4)

    p = doc.add_paragraph(
        'leider konnten wir bis heute keinen Zahlungseingang für die oben genannte Rechnung feststellen. '
        'Wir bitten Sie, den ausstehenden Betrag umgehend zu begleichen.'
    )
    p.space_after = Pt(12)

    # Summary table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    data = [
        ('Rechnungsbetrag:', '{{rechnung_betrag}}'),
        ('Mahngebühr:', '{{mahngebuehr}}'),
        ('Zahlungsfrist:', '{{mahnung_frist}}'),
    ]

    for i, (label, value) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        table.rows[i].cells[1].text = value
        table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    # Total row with highlight
    table.rows[3].cells[0].text = 'Gesamtbetrag:'
    table.rows[3].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_shading(table.rows[3].cells[0], 'FFF3F0')
    set_cell_shading(table.rows[3].cells[1], 'FFF3F0')

    # We cannot compute total in template, so just use invoice amount + fee placeholder
    table.rows[3].cells[1].text = '{{rechnung_betrag}} + {{mahngebuehr}}'
    table.rows[3].cells[1].paragraphs[0].runs[0].bold = True
    table.rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    table.rows[3].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xE1, 0x70, 0x55)

    doc.add_paragraph()

    p = doc.add_paragraph(
        'Bitte überweisen Sie den Gesamtbetrag bis zum {{mahnung_frist}} auf das unten genannte Konto. '
        'Sollte sich Ihre Zahlung mit diesem Schreiben überschnitten haben, betrachten Sie diese Mahnung bitte als gegenstandslos.'
    )
    p.space_after = Pt(8)

    p = doc.add_paragraph('{{notizen}}')
    p.space_after = Pt(8)

    p = doc.add_paragraph('Mit freundlichen Grüßen')
    p.space_after = Pt(4)
    p = doc.add_paragraph('{{inhaber}}')

    add_bank_footer(doc)

    doc.save(os.path.join('vorlagen', 'vorlage_mahnung.docx'))
    print('  vorlage_mahnung.docx erstellt')


if __name__ == '__main__':
    os.makedirs('vorlagen', exist_ok=True)
    print('Erstelle Word-Vorlagen...')
    create_invoice()
    create_offer()
    create_credit()
    create_reminder()
    print('Fertig! Alle Vorlagen liegen im Ordner "vorlagen/".')
